# MAIN.PY
import base64
import html
import math
import os
import re
import tempfile
import time
import uuid
from contextlib import contextmanager
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace
from typing import Any, Dict, List, Optional, Union

import psycopg2.pool
from fastapi import File, Form, HTTPException, UploadFile, status
from fastapi import FastAPI
from fastapi.responses import HTMLResponse
from openai import OpenAI
from pydantic import BaseModel, Field
from sentence_transformers import SentenceTransformer

# LangChain imports
from langchain_core.documents import Document
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnableParallel, RunnablePassthrough
from langchain_core.output_parsers import StrOutputParser
from langchain_openai import ChatOpenAI

try:
    from langchain.text_splitter import RecursiveCharacterTextSplitter
except ImportError:  # pragma: no cover
    from langchain_text_splitters import RecursiveCharacterTextSplitter  # type: ignore

# imports settings from settings.py
from backend.logging_config import get_logger
from backend.settings import settings
from backend.utils import normalize_text, vector_to_list

logger = get_logger(__name__)

# define the structure of the incoming request json
class AttachmentPayload(BaseModel):
    filename: str
    data: str
    mime_type: Optional[str] = None


class QueryRequest(BaseModel):
    query: str
    user_groups: List[str]
    attachments: List[AttachmentPayload] = Field(default_factory=list)

# define the structure for a single source document
class DocumentSource(BaseModel):
    text: str
    source: str
    doc_id: Optional[str] = None

# define the structure of the final json response
class QueryResponse(BaseModel):
    response: str
    sources: List[DocumentSource]

# Globals initialized during FastAPI startup
embedding_model = None
vllm_client = None
langchain_llm = None

MAX_RESPONSE_SOURCES = int(os.getenv("MAX_RESPONSE_SOURCES", "5"))

_TEST_STUB_CHUNKS: List[Dict[str, Union[str, float]]] = [
    {
        "doc_id": "stub-doc-1",
        "text": "This is a stubbed chunk used during automated tests.",
        "source": "stub_document.md",
        "similarity": 0.99,
    }
]
def _decode_attachment_data(encoded: str) -> bytes:
    """Decode base64 attachment data, enforcing strict validation."""
    try:
        return base64.b64decode(encoded, validate=True)
    except Exception as exc:  # pragma: no cover - defensive
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Attachment payload must be valid base64 data",
        ) from exc


def validate_query_request(request: QueryRequest) -> None:
    """Validate query request parameters."""
    if not request.query or len(request.query.strip()) < 3:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST, 
            detail="Query must be at least 3 characters"
        )
    if len(request.query) > settings.max_query_length:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST, 
            detail=f"Query too long (max {settings.max_query_length} characters)"
        )
    if not request.user_groups:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST, 
            detail="At least one user group required"
        )

    allowed_extensions = {".pdf", ".docx", ".txt", ".md"}
    for attachment in request.attachments or []:
        extension = Path(attachment.filename).suffix.lower()
        if extension not in allowed_extensions:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=(
                    f"Unsupported attachment type for '{attachment.filename}'. "
                    "Allowed types: PDF, DOCX, TXT, MD"
                ),
            )

        raw_bytes = _decode_attachment_data(attachment.data)
        if not raw_bytes:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Attachment '{attachment.filename}' is empty",
            )
        if len(raw_bytes) > settings.max_file_size:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=(
                    f"Attachment '{attachment.filename}' exceeds size limit "
                    f"({settings.max_file_size // (1024 * 1024)}MB)"
                ),
            )

def validate_upload_file(file: UploadFile) -> None:
    """Validate uploaded file."""
    if file.size and file.size > settings.max_file_size:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE, 
            detail=f"File too large (max {settings.max_file_size // (1024*1024)}MB)"
        )

def sanitize_model_output(text: str) -> str:
    """Remove chain-of-thought blocks and basic unsafe HTML tags from model output.

    - Strips <think> ... </think> blocks entirely
    - Removes script/style/iframe/object/embed tags if present
    - Trims excessive surrounding whitespace
    """
    try:
        import re

        if not isinstance(text, str):
            return ""

        # Remove <think>...</think> blocks (non-greedy, dotall, case-insensitive)
        text = re.sub(r"<\s*think\b[\s\S]*?>[\s\S]*?<\s*/\s*think\s*>", "", text, flags=re.IGNORECASE)
        # Also remove lone <think> or </think> tags if any remain
        text = re.sub(r"<\/?\s*think\s*>", "", text, flags=re.IGNORECASE)

        # Drop dangerous tags entirely (keep inner text out to be safe)
        for tag in ["script", "style", "iframe", "object", "embed"]:
            pattern = rf"<\s*{tag}\b[\s\S]*?>[\s\S]*?<\s*/\s*{tag}\s*>"
            text = re.sub(pattern, "", text, flags=re.IGNORECASE)

        # Basic cleanup of stray HTML tags that might affect rendering
        # Keep markdown; only remove remaining HTML tags conservatively
        text = re.sub(r"</?\w+[^>]*>", "", text)

        return text.strip()
    except Exception:
        # Fail-closed: return original text rather than crashing request
        return (text or "").strip()

def get_relevant_chunks_langchain(query: str, user_groups: List[str], limit: int = 25, similarity_threshold: float = 0.7) -> List[Document]:
    """
    LangChain-compatible retriever that returns Document objects.
    Uses the same SQL logic as get_relevant_chunks but returns LangChain Documents.
    """
    if settings.test_mode:
        logger.debug("Test mode active; returning stubbed LangChain documents")
        return [
            Document(
                page_content=chunk["text"],
                metadata={
                    "source": chunk["source"],
                    "similarity": chunk["similarity"],
                    "doc_id": chunk["doc_id"],
                },
            )
            for chunk in _TEST_STUB_CHUNKS[: min(limit, len(_TEST_STUB_CHUNKS))]
        ]

    try:
        # Generate embedding for the query
        query_embedding = vector_to_list(embedding_model.encode(query))
        
        with get_db_connection() as conn:
            cursor = conn.cursor()
            
            logger.info(f"LangChain: Searching for user with groups: {user_groups}")
            
            # Convert embedding to postgresql array format
            embedding_array = "[" + ",".join(map(str, query_embedding)) + "]"
            
            # SQL query with similarity threshold
            query_sql = """
            SELECT doc_id, chunk_text, source_path, 
                   1 - (embedding <=> %s::vector) as similarity
            FROM vector_index 
            WHERE allowed_groups && %s::text[]
            AND (1 - (embedding <=> %s::vector)) >= %s
            ORDER BY embedding <=> %s::vector
            LIMIT %s
            """
            
            fetch_limit = min(limit * 3, 100)
            cursor.execute(query_sql, (embedding_array, user_groups, embedding_array, similarity_threshold, embedding_array, fetch_limit))
            raw_results = cursor.fetchall()
            
            # Convert to LangChain Documents
            documents = []
            seen_ids = set()
            for row in raw_results:
                doc_id = row[0]
                if doc_id in seen_ids:
                    continue
                seen_ids.add(doc_id)
                doc = Document(
                    page_content=row[1],
                    metadata={
                        "source": row[2],
                        "similarity": float(row[3]),
                        "doc_id": doc_id,
                    }
                )
                documents.append(doc)
            
            logger.info(f"LangChain: Retrieved {len(documents)} relevant documents")
            cursor.close()
            return documents

    except Exception as e:
        logger.error(f"LangChain database error: {e}")
        return []

# LangChain prompt template
RAG_PROMPT = ChatPromptTemplate.from_messages([
    ("system", """You are Scoop, the internal AI assistant for Blue Bell Creameries.

Guidelines:
- Help employees get accurate answers with a friendly professional tone.
- The context will be provided as a numbered list of documents. Incorporate every relevant document; if multiple records refer to distinct items, mention each one explicitly.
- Use the provided document context when it is relevant and cite sources.
- Cite at most three sources. Only cite a source if you explicitly reference its information.
- If the context does not answer the question, say so and provide a concise general response if you know it.
- If you are unsure or lack knowledge, explicitly state that you do not know rather than guessing.
- Never fabricate citations or details; avoid hallucinations.
- Do not reveal internal reasoning or chain-of-thought—share only the final answer and brief citations when used."""),
    ("user", """Context from company documents:
{context}

Question: {question}

Please provide a helpful answer based on the context above.""")
])

from backend.utils import normalize_text

# create the FastAPI application
app = FastAPI(
    title="AI Assistant API",
    description="API for the AI assistant.",
    version="0.1.0",
)

# Database connection pool (initialized on startup)
db_pool = None

@app.on_event("startup")
def on_startup() -> None:
    """Initialize heavy resources: embeddings model, DB pool, LLM clients."""
    global embedding_model, db_pool, vllm_client, langchain_llm
    if settings.test_mode:
        logger.info("Test mode detected; installing lightweight stubs")

        # Minimal embedding model for deterministic behaviour in tests
        class _StubEmbeddingModel:
            def encode(self, value: Union[str, List[str]]) -> List[float]:
                if isinstance(value, list):
                    return [float(len(" ".join(value)) or 1.0)]
                if not isinstance(value, str):
                    return [0.0]
                return [float(len(value) or 1.0)]

        embedding_model = _StubEmbeddingModel()

        def _stub_completion_create(*args, **kwargs):
            return SimpleNamespace(
                choices=[
                    SimpleNamespace(
                        message=SimpleNamespace(
                            content="Stub response generated during test mode."
                        )
                    )
                ]
            )

        vllm_client = SimpleNamespace(
            chat=SimpleNamespace(
                completions=SimpleNamespace(create=_stub_completion_create)
            )
        )
        langchain_llm = SimpleNamespace(
            invoke=lambda *args, **kwargs: "Stub response generated during test mode."
        )
        db_pool = None
        return

    logger.info("Loading embedding model...")
    embedding_model = SentenceTransformer(settings.embedding_model_name)
    logger.info("Embedding model loaded successfully")

    logger.info("Initializing database connection pool...")
    db_pool = psycopg2.pool.SimpleConnectionPool(
        minconn=1,
        maxconn=20,
        host=settings.db_host,
        database=settings.db_name,
        user=settings.db_user,
        password=settings.db_password.get_secret_value(),
    )
    logger.info("Database connection pool initialized successfully")

    logger.info("Initializing vLLM client...")
    vllm_client = OpenAI(
        base_url=settings.vllm_base_url,
        api_key=settings.vllm_api_key.get_secret_value()
        if settings.vllm_api_key
        else "dummy-key",
    )
    logger.info("vLLM client initialized successfully")

    logger.info("Initializing LangChain ChatOpenAI client...")
    langchain_llm = ChatOpenAI(
        base_url=settings.vllm_base_url,
        api_key=settings.vllm_api_key.get_secret_value()
        if settings.vllm_api_key
        else "dummy-key",
        model=settings.vllm_model,
        temperature=0.7,
        max_tokens=1000,
    )
    logger.info("LangChain ChatOpenAI client initialized successfully")

@app.on_event("shutdown")
def on_shutdown() -> None:
    """Tear down heavy resources cleanly."""
    global db_pool
    try:
        if db_pool is not None:
            db_pool.closeall()
            logger.info("Database connection pool closed")
    except Exception as e:
        logger.warning(f"Error closing DB pool: {e}")

@contextmanager
def get_db_connection():
    """
    a helper function to get a connection from the pool and ensure
    it's returned, even if an error occurs.
    """
    if db_pool is None:
        raise RuntimeError("Database connection pool is not initialized")
    conn = db_pool.getconn()
    try:
        yield conn
    finally:
        db_pool.putconn(conn)


def get_relevant_chunks(query_embedding: List[float], user_groups: List[str], limit: int = 25, similarity_threshold: float = 0.7) -> List[Dict[str, Union[str, float]]]:
    """
    connects to the database via the pool and retrieves chunks using vector similarity search.
    """
    if settings.test_mode:
        logger.debug("Test mode active; returning stubbed chunk results")
        return _TEST_STUB_CHUNKS[: min(limit, len(_TEST_STUB_CHUNKS))]

    try:
        with get_db_connection() as conn:  # borrows a connection from the pool
            cursor = conn.cursor()
            
            logger.info(f"Searching for user with groups: {user_groups}")
            
            # convert embedding to postgresql array format
            embedding_array = "[" + ",".join(map(str, query_embedding)) + "]"
            
            # create user groups filter for permission-based search
            # (groups are passed as array param; no string concat needed)
            
            # first, get all matching documents with similarity scores for debugging
            debug_query = """
            SELECT doc_id, chunk_text, source_path, 
                   1 - (embedding <=> %s::vector) as similarity
            FROM vector_index 
            WHERE allowed_groups && %s::text[]
            ORDER BY embedding <=> %s::vector
            LIMIT 10
            """
            
            cursor.execute(debug_query, (embedding_array, user_groups, embedding_array))
            debug_results = cursor.fetchall()
            
            logger.info(f"Debug - All matching documents and their similarity scores:")
            for row in debug_results:
                logger.info(f"  - {row[2]} ({row[0]}): {row[3]:.4f}")
            
            # now apply similarity threshold
            query = """
            SELECT doc_id, chunk_text, source_path, 
                   1 - (embedding <=> %s::vector) as similarity
            FROM vector_index 
            WHERE allowed_groups && %s::text[]
            AND (1 - (embedding <=> %s::vector)) >= %s
            ORDER BY embedding <=> %s::vector
            LIMIT %s
            """
            
            fetch_limit = min(limit * 3, 100)
            cursor.execute(query, (embedding_array, user_groups, embedding_array, similarity_threshold, embedding_array, fetch_limit))
            raw_results = cursor.fetchall()
            
            # convert results to chunks
            chunks = []
            seen_ids = set()
            for row in raw_results:
                doc_id = row[0]
                if doc_id in seen_ids:
                    continue
                seen_ids.add(doc_id)
                chunks.append({
                    "doc_id": doc_id,
                    "text": row[1],
                    "source": row[2],
                    "similarity": float(row[3])
                })
            
            logger.info(f"Retrieved {len(chunks)} relevant chunks (similarity >= {similarity_threshold})")
            cursor.close()
            return chunks

    except Exception as e:
        logger.error(f"Database error: {e}")
        return []


def generate_response_with_vllm(query: str, context_chunks: List[Dict[str, Union[str, float]]]) -> str:
    """
    Generate a response using vLLM based on the query and retrieved context chunks.
    """
    if settings.test_mode:
        return "Stub response generated during test mode."

    try:
        # prepare context from retrieved chunks
        if not context_chunks:
            context = "No relevant documents found."
        else:
            context_entries = []
            for idx, chunk in enumerate(context_chunks, start=1):
                doc_id = chunk.get("doc_id", f"doc{idx}")
                context_entries.append(
                    f"Document {idx} (doc_id={doc_id}, source={chunk['source']}):\n{chunk['text']}"
                )
            context = "\n\n".join(context_entries)
        
        # create system prompt for RAG
        system_prompt = """You are Scoop, the internal AI assistant for Blue Bell Creameries.

Guidelines:
- Help employees with accurate, concise answers in a friendly tone.
- Prefer the provided context and cite sources when you use it.
- Cite at most three sources. Only cite a source if you explicitly use it in your answer.
- If the context does not contain the answer, say so and provide a careful general response if you know it.
- If you are unsure or information is unavailable, admit it rather than guessing.
- Never fabricate details or citations."""
        
        # prepare the user message with context
        user_message = f"""Context from company documents:
{context}

Question: {query}

Please provide a helpful answer based on the context above."""
        
        # make request to vLLM
        request_id = str(uuid.uuid4())[:8]
        logger.info(f"[{request_id}] Making vLLM request for query: {query[:50]}...")
        start_time = time.time()
        response = vllm_client.chat.completions.create(
            model=settings.vllm_model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message}
            ],
            temperature=0.7,
            max_tokens=1000,
            top_p=0.9
        )
        response_time = time.time() - start_time
        logger.info(f"[{request_id}] vLLM response generated in {response_time:.2f}s")
        
        return response.choices[0].message.content
        
    except Exception as e:
        logger.error(f"Error generating response with vLLM: {e}")
        return f"I apologize, but I encountered an error while generating a response: {str(e)}"


@app.get("/health", tags=["System"])
def health_check():
    """checks if the API is running"""
    return {"status": "ok"}

@app.get("/health/vllm", tags=["System"])
def vllm_health_check():
    """checks if vLLM is accessible"""
    if settings.test_mode:
        return {"status": "ok", "vllm": "stub"}
    try:
        # simple test request to vLLM
        response = vllm_client.chat.completions.create(
            model=settings.vllm_model,
            messages=[{"role": "user", "content": "test"}],
            max_tokens=1
        )
        return {"status": "ok", "vllm": "accessible"}
    except Exception as e:
        return {"status": "error", "vllm": "unavailable", "error": str(e)}


def process_query_common(request: QueryRequest, limit: int = 3, similarity_threshold: float = 0.3) -> QueryResponse:
    """
    Common query processing logic used by all query endpoints.
    """
    try:
        # Validate request
        validate_query_request(request)

        if request.attachments:
            logger.info(
                "Received %s transient attachment(s): %s",
                len(request.attachments),
                ", ".join(att.filename for att in request.attachments),
            )

        # embed the user's query using embedding model
        logger.info(f"Processing query: {request.query}")
        query_embedding = vector_to_list(embedding_model.encode(request.query))
        logger.info(f"Generated embedding with {len(query_embedding)} dimensions")

        # retrieve relevant chunks from the database
        chunks = get_relevant_chunks(query_embedding, request.user_groups, limit=limit, similarity_threshold=similarity_threshold)

        attachment_chunks = build_ephemeral_attachment_chunks(query_embedding, request.attachments)
        if attachment_chunks:
            logger.info(
                "Including %s transient chunk(s) from attachments",
                len(attachment_chunks),
            )
            chunks = attachment_chunks + chunks

        # convert chunks to DocumentSource format
        sources = select_response_sources(chunks)
        
        # generate response using vLLM
        llm_response = generate_response_with_vllm(request.query, chunks)
        llm_response = sanitize_model_output(llm_response)
        
        logger.info(f"Returning response with {len(sources)} sources")
        return {"response": llm_response, "sources": sources}
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error processing query: {e}")
        raise HTTPException(status_code=500, detail=f"Error processing query: {str(e)}")

@app.post("/query", tags=["RAG"], response_model=QueryResponse)
def process_query(request: QueryRequest):
    """
    takes user query and user groups, retrieves relevant context,
    and returns an answer.
    """
    return process_query_common(request, limit=25, similarity_threshold=0.3)


@app.post("/query-precise", tags=["RAG"], response_model=QueryResponse)
def process_query_precise(request: QueryRequest):
    """
    takes user query and user groups, retrieves only highly relevant context,
    and returns an answer with stricter similarity filtering.
    """
    return process_query_common(request, limit=12, similarity_threshold=0.5)


@app.post("/query-lc", tags=["RAG"], response_model=QueryResponse)
def process_query_langchain(request: QueryRequest):
    """
    LangChain-powered query processing using LCEL (LangChain Expression Language).
    Uses the same retriever logic but with LangChain's chain composition.
    """
    if settings.test_mode:
        logger.info("Test mode: returning stub LangChain response")
        sources = [
            DocumentSource(
                text=chunk["text"],
                source=f"{chunk['source']} ({chunk['doc_id']})",
                doc_id=chunk["doc_id"],
            )
            for chunk in _TEST_STUB_CHUNKS
        ]
        return {
            "response": "Stub LangChain response generated during test mode.",
            "sources": sources,
        }

    try:
        logger.info(f"Processing LangChain query: {request.query}")
        
        query_embedding = vector_to_list(embedding_model.encode(request.query))

        attachment_chunks = build_ephemeral_attachment_chunks(
            query_embedding,
            request.attachments,
        )

        if attachment_chunks:
            logger.info(
                "LangChain flow including %s transient chunk(s) from attachments",
                len(attachment_chunks),
            )

        # Retrieve documents using LangChain-compatible retriever
        documents = get_relevant_chunks_langchain(
            request.query,
            request.user_groups,
            limit=25,
            similarity_threshold=0.3,
        )

        if attachment_chunks:
            documents = [
                Document(
                    page_content=chunk["text"],
                    metadata={
                        "source": chunk["source"],
                        "similarity": chunk["similarity"],
                        "doc_id": chunk["doc_id"],
                    },
                )
                for chunk in attachment_chunks
            ] + documents
        
        if not documents:
            # Fallback to a generic assistant answer without context
            logger.info("No documents found – falling back to general assistant response")
            try:
                raw = langchain_llm.invoke(
                    "You are Scoop, the internal AI assistant for Blue Bell Creameries.\n\n"
                    "Provide accurate, concise answers. If you lack information, say so instead of guessing.\n\n"
                    f"Question: {request.query}\n\nAnswer helpfully."
                )
                # Extract text if ChatMessage
                if hasattr(raw, "content"):
                    raw = raw.content
                text = str(raw).strip()
                cleaned = sanitize_model_output(text)
                if not cleaned:
                    cleaned = text  # use raw if sanitizer stripped too much
                return {"response": cleaned, "sources": []}
            except Exception as e:
                logger.error(f"Fallback failed: {e}")
                return {"response": "I'm sorry, I couldn't answer that.", "sources": []}
        
        # Format context from documents
        context = "\n\n".join([
            f"Document {idx} (doc_id={doc.metadata.get('doc_id', f'doc{idx}')}, source={doc.metadata['source']}):\n{doc.page_content}"
            for idx, doc in enumerate(documents, start=1)
        ])
        
        # Create the LangChain chain
        chain = (
            RunnableParallel({
                "context": lambda x: context,
                "question": RunnablePassthrough()
            })
            | RAG_PROMPT
            | langchain_llm
            | StrOutputParser()
        )
        
        # Invoke the chain
        request_id = str(uuid.uuid4())[:8]
        logger.info(f"[{request_id}] Invoking LangChain chain...")
        start_time = time.time()
        
        response_text = chain.invoke({"question": request.query})
        
        response_time = time.time() - start_time
        logger.info(f"[{request_id}] LangChain response generated in {response_time:.2f}s")
        
        # Sanitize the response
        response_text = sanitize_model_output(response_text)
        
        # Convert documents to DocumentSource format
        source_chunks = documents_to_chunks(documents)
        sources = select_response_sources(source_chunks)
        
        logger.info(f"Returning LangChain response with {len(sources)} sources")
        return {"response": response_text, "sources": sources}
        
    except Exception as e:
        logger.error(f"Error processing LangChain query: {e}")
        raise HTTPException(status_code=500, detail=f"Error processing query: {str(e)}")


def parse_document_content(file_path: Path) -> str:
    """Parse document content based on file extension."""
    ext = file_path.suffix.lower()
    
    if ext == '.txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    elif ext == '.md':
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import PyPDF2
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text.strip()
        except ImportError:
            raise HTTPException(status_code=400, detail="PDF parsing requires PyPDF2. Install with: pip install PyPDF2")
    elif ext == '.docx':
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except ImportError:
            raise HTTPException(status_code=400, detail="DOCX parsing requires python-docx. Install with: pip install python-docx")
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {ext}")
def chunk_text(text: str, chunk_size: int = 512, overlap: int = 50) -> List[str]:
    """Chunk text using RecursiveCharacterTextSplitter with smart separators."""

    if not text.strip():
        return []

    approx_chars_per_word = 5
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=max(chunk_size, 50) * approx_chars_per_word,
        chunk_overlap=max(overlap, 0) * approx_chars_per_word,
        separators=["\n\n", "\n", " ", ""],
    )

    chunks = [segment.strip() for segment in splitter.split_text(text) if segment.strip()]
    return chunks


def _extract_attachment_text(filename: str, raw_bytes: bytes) -> str:
    """Read attachment bytes and return normalized text content."""
    suffix = Path(filename).suffix.lower()

    if suffix in {".txt", ".md"}:
        return normalize_text(raw_bytes.decode("utf-8", errors="ignore"))

    if suffix == ".pdf":
        try:
            import PyPDF2
        except ImportError as exc:  # pragma: no cover - optional dependency
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="PDF support requires PyPDF2. Install with `pip install PyPDF2`.",
            ) from exc
        reader = PyPDF2.PdfReader(BytesIO(raw_bytes))
        extracted = "".join(page.extract_text() or "" for page in reader.pages)
        return normalize_text(extracted)

    if suffix == ".docx":
        try:
            from docx import Document as DocxDocument
        except ImportError as exc:  # pragma: no cover - optional dependency
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="DOCX support requires python-docx. Install with `pip install python-docx`.",
            ) from exc
        document = DocxDocument(BytesIO(raw_bytes))
        extracted = "\n".join(paragraph.text for paragraph in document.paragraphs)
        return normalize_text(extracted)

    raise HTTPException(
        status_code=status.HTTP_400_BAD_REQUEST,
        detail=f"Unsupported attachment type: {suffix or 'unknown'}",
    )


def _cosine_similarity(vec_a: List[float], vec_b: List[float]) -> float:
    """Compute cosine similarity between two dense vectors."""
    if not vec_a or not vec_b or len(vec_a) != len(vec_b):
        return 0.0

    dot_product = 0.0
    norm_a = 0.0
    norm_b = 0.0
    for a, b in zip(vec_a, vec_b):
        dot_product += a * b
        norm_a += a * a
        norm_b += b * b

    if norm_a <= 0 or norm_b <= 0:
        return 0.0

    return dot_product / (math.sqrt(norm_a) * math.sqrt(norm_b))


def build_ephemeral_attachment_chunks(
    query_embedding: List[float],
    attachments: List[AttachmentPayload],
    chunk_limit: int = 8,
) -> List[Dict[str, Union[str, float]]]:
    """Convert transient attachments into ranked chunks for a single response."""
    if not attachments:
        return []

    ranked_chunks: List[Dict[str, Union[str, float]]] = []

    for attachment_index, attachment in enumerate(attachments, start=1):
        raw_bytes = _decode_attachment_data(attachment.data)
        try:
            attachment_text = _extract_attachment_text(attachment.filename, raw_bytes)
        except HTTPException:
            raise
        except Exception as exc:  # pragma: no cover - defensive
            logger.error(
                "Failed to extract text from attachment '%s': %s",
                attachment.filename,
                exc,
            )
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Failed to read attachment '{attachment.filename}'",
            ) from exc

        if not attachment_text.strip():
            logger.info(
                "Skipping attachment '%s' because no readable text was found",
                attachment.filename,
            )
            continue

        chunks = chunk_text(
            attachment_text,
            chunk_size=settings.default_chunk_size,
            overlap=settings.default_chunk_overlap,
        )

        if not chunks:
            chunks = [attachment_text.strip()]

        for chunk_index, chunk_body in enumerate(chunks, start=1):
            chunk_header = (
                f"[Attachment {attachment_index} - Chunk {chunk_index}/{len(chunks)}]"
            )
            contextual_chunk = f"{chunk_header}\n{chunk_body}"
            embedding = vector_to_list(embedding_model.encode(contextual_chunk))
            similarity = _cosine_similarity(query_embedding, embedding)

            ranked_chunks.append(
                {
                    "doc_id": f"attachment-{attachment_index}-{chunk_index}",
                    "text": contextual_chunk,
                    "source": f"{attachment.filename} (attachment)",
                    "similarity": similarity,
                }
            )

    ranked_chunks.sort(key=lambda entry: entry.get("similarity", 0.0), reverse=True)
    if chunk_limit > 0:
        ranked_chunks = ranked_chunks[:chunk_limit]

    return ranked_chunks


def _chunk_is_attachment(chunk: Dict[str, Union[str, float]]) -> bool:
    doc_id = chunk.get("doc_id")
    return isinstance(doc_id, str) and doc_id.startswith("attachment-")


def _format_document_source(chunk: Dict[str, Union[str, float]]) -> DocumentSource:
    doc_id = chunk.get("doc_id")
    source_label = chunk.get("source", "unknown")
    if doc_id:
        source_label = f"{source_label} ({doc_id})"
    return DocumentSource(
        text=chunk.get("text", ""),
        source=source_label,
        doc_id=doc_id,
    )


def select_response_sources(
    chunks: List[Dict[str, Union[str, float]]],
    max_non_attachment_sources: int = MAX_RESPONSE_SOURCES,
) -> List[DocumentSource]:
    """
    Build a compact source list prioritizing attachments and the highest-similarity chunks.
    """
    if not chunks:
        return []

    attachments: List[Dict[str, Union[str, float]]] = []
    others: List[Dict[str, Union[str, float]]] = []
    seen_ids: set[str] = set()

    for chunk in chunks:
        doc_id = chunk.get("doc_id")
        if doc_id and doc_id in seen_ids:
            continue
        if isinstance(doc_id, str):
            seen_ids.add(doc_id)

        if _chunk_is_attachment(chunk):
            attachments.append(chunk)
        else:
            others.append(chunk)

    attachments.sort(key=lambda entry: entry.get("similarity", 0.0), reverse=True)
    others.sort(key=lambda entry: entry.get("similarity", 0.0), reverse=True)

    selected_chunks: List[Dict[str, Union[str, float]]] = []
    selected_chunks.extend(attachments)

    if max_non_attachment_sources > 0:
        remaining = max_non_attachment_sources
        for chunk in others:
            if remaining <= 0:
                break
            selected_chunks.append(chunk)
            remaining -= 1

    return [_format_document_source(chunk) for chunk in selected_chunks]


def documents_to_chunks(documents: List[Document]) -> List[Dict[str, Union[str, float]]]:
    """Convert LangChain Document objects to the chunk dict format."""
    chunk_dicts: List[Dict[str, Union[str, float]]] = []
    for doc in documents:
        metadata = getattr(doc, "metadata", {}) or {}
        chunk_dicts.append(
            {
                "doc_id": metadata.get("doc_id"),
                "text": doc.page_content,
                "source": metadata.get("source", "unknown"),
                "similarity": metadata.get("similarity", 0.0),
            }
        )
    return chunk_dicts


@app.post("/upload", tags=["Upload"])
async def upload_document(
    file: UploadFile = File(...),
    department: str = Form(...),
    allowed_groups: str = Form(...),
    chunk_size: int = Form(512),
    chunk_overlap: int = Form(50)
):
    """
    Upload and ingest a document with metadata.
    """
    try:
        # Validate file
        validate_upload_file(file)
        
        # Validate file type
        allowed_extensions = ['.txt', '.md', '.pdf', '.docx']
        file_ext = Path(file.filename).suffix.lower()
        
        if file_ext not in allowed_extensions:
            raise HTTPException(
                status_code=400, 
                detail=f"Unsupported file type. Allowed: {', '.join(allowed_extensions)}"
            )
        
        # Parse allowed groups
        groups = [g.strip() for g in allowed_groups.split(',') if g.strip()]
        
        # Save uploaded file temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as tmp_file:
            content = await file.read()
            tmp_file.write(content)
            tmp_path = Path(tmp_file.name)
        
        try:
            # Parse document content
            logger.info(f"Parsing document: {file.filename}")
            doc_content = parse_document_content(tmp_path)
            doc_content = normalize_text(doc_content)

            if not doc_content.strip():
                raise HTTPException(status_code=400, detail="Document appears to be empty")
            
            # Chunk the document
            chunks = chunk_text(doc_content, chunk_size, chunk_overlap)
            logger.info(f"Created {len(chunks)} chunks from {file.filename}")

            if chunks:
                preview = chunks[0]
                logger.info(
                    "Sample chunk [1/%s]: %s",
                    len(chunks),
                    preview[:200].replace("\n", " ") + ("..." if len(preview) > 200 else ""),
                )

            if not chunks:
                raise HTTPException(status_code=400, detail="Failed to produce any text chunks")

            # Process each chunk
            inserted_chunks = 0
            document_id = str(uuid.uuid4())
            try:
                with get_db_connection() as conn:
                    cursor = conn.cursor()
                    try:
                        total_chunks = len(chunks)
                        for i, raw_chunk in enumerate(chunks):
                            chunk_content = f"[Chunk {i + 1}/{total_chunks}]\n{raw_chunk}"

                            embedding = vector_to_list(embedding_model.encode(chunk_content))

                            cursor.execute(
                                """
                                INSERT INTO vector_index (
                                    doc_id, source_path, department, security_level,
                                    allowed_groups, last_updated_by, last_updated_at,
                                    chunk_text, embedding
                                )
                                VALUES (%s, %s, %s, %s, %s, %s, NOW(), %s, %s)
                                """,
                                (
                                    document_id,
                                    file.filename,
                                    department,
                                    "Internal",  # Default security_level
                                    groups,
                                    "web_upload",
                                    chunk_content,
                                    embedding,
                                ),
                            )
                            inserted_chunks += 1

                        conn.commit()
                    except Exception:
                        conn.rollback()
                        raise
                    finally:
                        cursor.close()
            except Exception as e:
                logger.error(f"Failed to insert chunks for {file.filename}: {e}")
                raise

            return {
                "message": f"Successfully uploaded and processed {file.filename}",
                "filename": file.filename,
                "chunks_created": len(chunks),
                "chunks_inserted": inserted_chunks,
                "department": department,
                "allowed_groups": groups
            }
            
        finally:
            # Clean up temporary file
            if tmp_path.exists():
                os.unlink(tmp_path)
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error uploading document: {e}")
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")


@app.get("/database", response_class=HTMLResponse, tags=["Database"])
def view_documents():
    """View all documents in the database."""
    try:
        with get_db_connection() as conn:
            cursor = conn.cursor()
            
            # Get document summary
            cursor.execute("""
                SELECT 
                    source_path,
                    department,
                    allowed_groups,
                    COUNT(*) as chunk_count,
                    MAX(last_updated_at) as last_updated
                FROM vector_index 
                GROUP BY source_path, department, allowed_groups
                ORDER BY last_updated DESC
            """)
            
            documents = cursor.fetchall()
            
            # Get total stats
            cursor.execute("SELECT COUNT(*) FROM vector_index")
            total_chunks = cursor.fetchone()[0]
            
            cursor.execute("SELECT COUNT(DISTINCT source_path) FROM vector_index")
            total_documents = cursor.fetchone()[0]
            
            cursor.close()
        
        # Build HTML
        page_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>Document Database - RAG System</title>
            <style>
                body {{ font-family: Arial, sans-serif; max-width: 1200px; margin: 0 auto; padding: 20px; }}
                .header {{ background-color: #f8f9fa; padding: 20px; border-radius: 8px; margin-bottom: 20px; }}
                .stats {{ display: flex; gap: 20px; margin-bottom: 20px; }}
                .stat-box {{ background-color: #e9ecef; padding: 15px; border-radius: 8px; text-align: center; min-width: 120px; }}
                .stat-number {{ font-size: 24px; font-weight: bold; color: #007bff; }}
                .stat-label {{ font-size: 14px; color: #6c757d; }}
                .documents-table {{ width: 100%; border-collapse: collapse; margin-top: 20px; }}
                .documents-table th, .documents-table td {{ border: 1px solid #dee2e6; padding: 12px; text-align: left; }}
                .documents-table th {{ background-color: #f8f9fa; font-weight: bold; }}
                .documents-table tr:nth-child(even) {{ background-color: #f8f9fa; }}
                .btn {{ padding: 8px 16px; border: none; border-radius: 4px; cursor: pointer; text-decoration: none; display: inline-block; margin: 2px; }}
                .btn-primary {{ background-color: #007bff; color: white; }}
                .btn-danger {{ background-color: #dc3545; color: white; }}
                .btn-secondary {{ background-color: #6c757d; color: white; }}
                .btn:hover {{ opacity: 0.8; }}
                .groups {{ font-size: 12px; color: #6c757d; }}
                .actions {{ white-space: nowrap; }}
            </style>
        </head>
        <body>
            <div class="header">
                <h1>Document Database</h1>
                <p>View and manage all documents in the RAG system</p>
            </div>
            
            <div class="stats">
                <div class="stat-box">
                    <div class="stat-number">{total_documents}</div>
                    <div class="stat-label">Documents</div>
                </div>
                <div class="stat-box">
                    <div class="stat-number">{total_chunks}</div>
                    <div class="stat-label">Chunks</div>
                </div>
            </div>
            
            <div style="margin-bottom: 20px;">
                <a href="/upload" class="btn btn-primary">📄 Upload New Document</a>
                <a href="/database" class="btn btn-secondary">🔄 Refresh</a>
            </div>
            
            <table class="documents-table">
                <thead>
                    <tr>
                        <th>Document</th>
                        <th>Department</th>
                        <th>Allowed Groups</th>
                        <th>Chunks</th>
                        <th>Last Updated</th>
                        <th>Actions</th>
                    </tr>
                </thead>
                <tbody>
        """
        
        for doc in documents:
            source_path, department, allowed_groups, chunk_count, last_updated = doc
            
            # Safely format the data (escape HTML to prevent XSS)
            safe_source_path = html.escape(str(source_path))
            safe_department = html.escape(str(department)) if department else "Unknown"
            
            # Handle allowed_groups safely (escape HTML)
            if allowed_groups and isinstance(allowed_groups, list):
                groups_str = html.escape(', '.join(str(g) for g in allowed_groups))
            else:
                groups_str = "None"
            
            # Format date safely (escape HTML)
            if last_updated:
                try:
                    date_str = html.escape(last_updated.strftime('%Y-%m-%d %H:%M'))
                except:
                    date_str = html.escape(str(last_updated))
            else:
                date_str = "Unknown"
            
            page_html += f"""
                    <tr>
                        <td><strong>{safe_source_path}</strong></td>
                        <td>{safe_department}</td>
                        <td class="groups">{groups_str}</td>
                        <td>{chunk_count}</td>
                        <td>{date_str}</td>
                        <td class="actions">
                            <button class="btn btn-primary" onclick="previewDocument('{safe_source_path}')">👁️ Preview</button>
                            <button class="btn btn-danger" onclick="deleteDocument('{safe_source_path}')">🗑️ Delete</button>
                        </td>
                    </tr>
            """
        
        page_html += """
                </tbody>
            </table>
            
            <script>
                function previewDocument(filename) {
                    fetch(`/docs/preview/${encodeURIComponent(filename)}`)
                        .then(response => response.json())
                        .then(data => {
                            if (data.error) {
                                alert('Error: ' + data.error);
                            } else {
                                let content = 'Document: ' + data.filename + '\\n\\n';
                                content += 'Department: ' + data.department + '\\n';
                                content += 'Allowed Groups: ' + data.allowed_groups.join(', ') + '\\n';
                                content += 'Chunks: ' + data.chunks.length + '\\n\\n';
                                content += 'Content Preview:\\n';
                                content += data.chunks.slice(0, 2).map((chunk, i) => 
                                    `Chunk ${i+1}: ${chunk.substring(0, 200)}...`
                                ).join('\\n\\n');
                                alert(content);
                            }
                        })
                        .catch(error => {
                            alert('Error loading preview: ' + error.message);
                        });
                }
                
                function deleteDocument(filename) {
                    if (confirm('Are you sure you want to delete all chunks for "' + filename + '"? This cannot be undone.')) {
                        fetch(`/docs/delete/${encodeURIComponent(filename)}`, { method: 'DELETE' })
                            .then(response => response.json())
                            .then(data => {
                                if (data.success) {
                                    alert('Document deleted successfully!');
                                    location.reload();
                                } else {
                                    alert('Error: ' + data.error);
                                }
                            })
                            .catch(error => {
                                alert('Error deleting document: ' + error.message);
                            });
                    }
                }
            </script>
        </body>
        </html>
        """
        
        return page_html
        
    except Exception as e:
        logger.error(f"Error viewing documents: {e}")
        return f"<h1>Error</h1><p>Failed to load documents: {str(e)}</p>"


@app.get("/docs/preview/{filename}", tags=["Database"])
def preview_document(filename: str):
    """Preview a document's chunks."""
    try:
        with get_db_connection() as conn:
            cursor = conn.cursor()
            
            cursor.execute("""
                SELECT chunk_text, department, allowed_groups
                FROM vector_index 
                WHERE source_path = %s
                ORDER BY doc_id
            """, (filename,))
            
            results = cursor.fetchall()
            
            if not results:
                return {"error": "Document not found"}
            
            chunks = [row[0] for row in results]
            department = results[0][1]
            allowed_groups = results[0][2]
            
            cursor.close()
            
            return {
                "filename": filename,
                "department": department,
                "allowed_groups": allowed_groups,
                "chunks": chunks
            }
            
    except Exception as e:
        logger.error(f"Error previewing document: {e}")
        return {"error": str(e)}


@app.delete("/docs/delete/{filename}", tags=["Database"])
def delete_document(filename: str):
    """Delete all chunks for a document."""
    try:
        with get_db_connection() as conn:
            cursor = conn.cursor()
            
            cursor.execute("DELETE FROM vector_index WHERE source_path = %s", (filename,))
            deleted_count = cursor.rowcount
            
            conn.commit()
            cursor.close()
            
            return {
                "success": True,
                "message": f"Deleted {deleted_count} chunks for {filename}",
                "deleted_count": deleted_count
            }
            
    except Exception as e:
        logger.error(f"Error deleting document: {e}")
        return {"success": False, "error": str(e)}


@app.get("/upload", response_class=HTMLResponse, tags=["Upload"])
def upload_form():
    """Serve the upload form."""
    return """
    <!DOCTYPE html>
    <html>
    <head>
        <title>Document Upload - RAG System</title>
        <style>
            body { font-family: Arial, sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; }
            .form-group { margin-bottom: 20px; }
            label { display: block; margin-bottom: 5px; font-weight: bold; }
            input, select, textarea { width: 100%; padding: 10px; border: 1px solid #ddd; border-radius: 4px; }
            button { background-color: #007bff; color: white; padding: 12px 24px; border: none; border-radius: 4px; cursor: pointer; font-size: 16px; }
            button:hover { background-color: #0056b3; }
            .result { margin-top: 20px; padding: 15px; border-radius: 4px; }
            .success { background-color: #d4edda; border: 1px solid #c3e6cb; color: #155724; }
            .error { background-color: #f8d7da; border: 1px solid #f5c6cb; color: #721c24; }
        </style>
    </head>
    <body>
        <h1>Document Upload</h1>
        <p>Upload documents to add them to the RAG system. Select the appropriate metadata and access controls.</p>
        
        <form id="uploadForm" enctype="multipart/form-data">
            <div class="form-group">
                <label for="file">Choose Document:</label>
                <input type="file" id="file" name="file" accept=".txt,.md,.pdf,.docx" required>
                <small>Supported formats: TXT, MD, PDF, DOCX</small>
            </div>
            
            <div class="form-group">
                <label for="department">Department:</label>
                <select id="department" name="department" required>
                    <option value="">Select Department</option>
                    <option value="Finance">Finance</option>
                    <option value="Engineering">Engineering</option>
                    <option value="Human Resources">Human Resources</option>
                    <option value="Security">Security</option>
                    <option value="Product">Product</option>
                    <option value="Marketing">Marketing</option>
                    <option value="Sales">Sales</option>
                    <option value="Operations">Operations</option>
                    <option value="General">General</option>
                </select>
            </div>
            
            
            <div class="form-group">
                <label for="allowed_groups">Allowed Groups (comma-separated):</label>
                <input type="text" id="allowed_groups" name="allowed_groups" 
                       placeholder="e.g., finance,executives,management" required>
                <small>Examples: finance,executives | engineering,product | hr,employees,management</small>
            </div>
            
            <div class="form-group">
                <label for="chunk_size">Chunk Size (words):</label>
                <input type="number" id="chunk_size" name="chunk_size" value="512" min="100" max="2048">
                <small>Number of words per chunk (100-2048)</small>
            </div>
            
            <div class="form-group">
                <label for="chunk_overlap">Chunk Overlap (words):</label>
                <input type="number" id="chunk_overlap" name="chunk_overlap" value="50" min="0" max="200">
                <small>Overlap between chunks for context preservation</small>
            </div>
            
            <button type="submit">Upload & Ingest Document</button>
        </form>
        
        <div id="result"></div>
        
        <script>
            document.getElementById('uploadForm').addEventListener('submit', async function(e) {
                e.preventDefault();
                
                const formData = new FormData();
                formData.append('file', document.getElementById('file').files[0]);
                formData.append('department', document.getElementById('department').value);
                formData.append('allowed_groups', document.getElementById('allowed_groups').value);
                formData.append('chunk_size', document.getElementById('chunk_size').value);
                formData.append('chunk_overlap', document.getElementById('chunk_overlap').value);
                
                const resultDiv = document.getElementById('result');
                resultDiv.innerHTML = '<p>Uploading and processing document...</p>';
                
                try {
                    const response = await fetch('/upload', {
                        method: 'POST',
                        body: formData
                    });
                    
                    const data = await response.json();
                    
                    if (response.ok) {
                        resultDiv.innerHTML = `
                            <div class="result success">
                                <h3>Upload Successful!</h3>
                                <p><strong>File:</strong> ${data.filename}</p>
                                <p><strong>Chunks Created:</strong> ${data.chunks_created}</p>
                                <p><strong>Chunks Inserted:</strong> ${data.chunks_inserted}</p>
                                <p><strong>Department:</strong> ${data.department}</p>
                                <p><strong>Allowed Groups:</strong> ${data.allowed_groups.join(', ')}</p>
                            </div>
                        `;
                    } else {
                        resultDiv.innerHTML = `
                            <div class="result error">
                                <h3>Upload Failed</h3>
                                <p>${data.detail}</p>
                            </div>
                        `;
                    }
                } catch (error) {
                    resultDiv.innerHTML = `
                        <div class="result error">
                            <h3>Upload Error</h3>
                            <p>${error.message}</p>
                        </div>
                    `;
                }
            });
        </script>
    </body>
    </html>
    """